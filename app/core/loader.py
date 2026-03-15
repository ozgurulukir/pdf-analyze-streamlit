"""Document loader for various file types."""
import io
import os
from typing import List, Optional
import streamlit as st
from PyPDF2 import PdfReader

try:
    from langchain_core.documents import Document
except ImportError:
    from langchain.schema import Document


class DocumentLoader:
    """Handles loading various document formats."""

    SUPPORTED_TYPES = {
        "pdf": "PDF Document",
        "txt": "Text File",
        "md": "Markdown File",
        "docx": "Word Document",
        "html": "HTML File",
    }

    @staticmethod
    def load_pdf(file) -> str:
        """Load and extract text from PDF."""
        try:
            # Handle both BytesIO and UploadedFile
            if hasattr(file, 'read'):
                file.seek(0)
                pdf_reader = PdfReader(file)
            else:
                with open(file, 'rb') as f:
                    pdf_reader = PdfReader(f)
            
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            return text
        except Exception as e:
            st.error(f"PDF okuma hatası: {str(e)}")
            return ""

    @staticmethod
    def load_txt(file) -> str:
        """Load and extract text from text file."""
        try:
            # Try multiple encodings
            if hasattr(file, 'read'):
                file.seek(0)
                content = file.read()
                
                for encoding in ['utf-8', 'latin-1', 'cp1252', 'utf-16']:
                    try:
                        return content.decode(encoding)
                    except (UnicodeDecodeError, AttributeError):
                        continue
                
                # Last resort: decode with errors='ignore'
                return content.decode('utf-8', errors='ignore')
            else:
                with open(file, 'r', encoding='utf-8') as f:
                    return f.read()
        except Exception as e:
            st.error(f"TXT okuma hatası: {str(e)}")
            return ""

    @staticmethod
    def load_docx(file) -> str:
        """Load and extract text from DOCX."""
        try:
            from docx import Document as DocxDocument
            
            # Handle both BytesIO and file path
            if hasattr(file, 'read'):
                file.seek(0)
                doc = DocxDocument(io.BytesIO(file.read()))
            else:
                doc = DocxDocument(file)
            
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            return text
        except Exception as e:
            st.error(f"DOCX okuma hatası: {str(e)}")
            return ""

    @staticmethod
    def load_html(file) -> str:
        """Load and extract text from HTML."""
        try:
            from bs4 import BeautifulSoup
            
            if hasattr(file, 'read'):
                file.seek(0)
                content = file.read()
                soup = BeautifulSoup(content, 'lxml')
            else:
                with open(file, 'r', encoding='utf-8') as f:
                    soup = BeautifulSoup(f, 'lxml')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            return soup.get_text(separator='\n', strip=True)
        except Exception as e:
            st.error(f"HTML okuma hatası: {str(e)}")
            return ""

    @classmethod
    def load_file(cls, uploaded_file) -> str:
        """Load file based on extension."""
        file_ext = uploaded_file.name.split('.')[-1].lower()

        if file_ext not in cls.SUPPORTED_TYPES:
            raise ValueError(f"Desteklenmeyen dosya tipi: {file_ext}")

        if file_ext == "pdf":
            return cls.load_pdf(uploaded_file)
        elif file_ext in ("txt", "md"):
            return cls.load_txt(uploaded_file)
        elif file_ext == "docx":
            return cls.load_docx(uploaded_file)
        elif file_ext == "html":
            return cls.load_html(uploaded_file)

        return ""

    @classmethod
    def load_documents(cls, uploaded_files) -> List[Document]:
        """Load multiple documents and return as LangChain Documents."""
        documents = []

        for uploaded_file in uploaded_files:
            try:
                text = cls.load_file(uploaded_file)
                if text.strip():
                    doc = Document(
                        page_content=text,
                        metadata={
                            "source": uploaded_file.name,
                            "type": uploaded_file.name.split('.')[-1].lower(),
                            "size": uploaded_file.size,
                        }
                    )
                    documents.append(doc)
            except Exception as e:
                st.warning(f"⚠️ {uploaded_file.name} yüklenemedi: {str(e)}")

        return documents

    @classmethod
    def load_directory(cls, directory_path: str, glob_pattern: str = "**/*.pdf") -> List[Document]:
        """Load documents from a directory using DirectoryLoader."""
        from langchain_community.document_loaders import DirectoryLoader, PyPDFLoader, TextLoader, UnstructuredMarkdownLoader
        
        # Mapping extension to loader
        loaders = {
            ".pdf": PyPDFLoader,
            ".txt": TextLoader,
            ".md": UnstructuredMarkdownLoader,
        }
        
        documents = []
        path = os.path.abspath(directory_path)
        
        if not os.path.exists(path):
            st.error(f"Dizin bulunamadı: {path}")
            return []

        try:
            loader = DirectoryLoader(
                path,
                glob=glob_pattern,
                show_progress=True,
                use_multithreading=True
            )
            documents = loader.load()
            return documents
        except Exception as e:
            st.error(f"Dizin yükleme hatası: {str(e)}")
            return []

    @staticmethod
    def validate_file(uploaded_file, max_size_mb: int = 50) -> tuple[bool, Optional[str]]:
        """Validate file before processing."""
        file_ext = uploaded_file.name.split('.')[-1].lower()
        
        if file_ext not in DocumentLoader.SUPPORTED_TYPES:
            return False, f"Desteklenmeyen dosya tipi: {file_ext}"
        
        if uploaded_file.size > max_size_mb * 1024 * 1024:
            return False, f"Dosya boyutu çok büyük (max {max_size_mb}MB)"
        
        return True, None
